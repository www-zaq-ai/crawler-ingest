"""
docx_to_md.py — Convert .docx files to clean markdown.

Uses mammoth for text/headings/lists and python-docx for tables,
stitched together in document order.

Usage:
    # Single file
    python docx_to_md.py report.docx

    # Single file with custom output path
    python docx_to_md.py report.docx --output report.md

    # Entire folder
    python docx_to_md.py --input-folder ./docs --output-folder ./markdown

    # Quiet mode
    python docx_to_md.py report.docx --quiet
"""

import argparse
import re
import sys
from pathlib import Path
from typing import List, Optional, Set

try:
    import mammoth
except ImportError:
    print("mammoth is not installed. Run: pip install mammoth")
    sys.exit(1)

try:
    from docx import Document as DocxDocument
except ImportError:
    print("python-docx is not installed. Run: pip install python-docx")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Custom style map for mammoth
# ---------------------------------------------------------------------------
STYLE_MAP = """
p[style-name='Title'] => h1.title:fresh
p[style-name='Subtitle'] => p.subtitle:fresh
p[style-name='Heading 1'] => h1:fresh
p[style-name='Heading 2'] => h2:fresh
p[style-name='Heading 3'] => h3:fresh
p[style-name='Heading 4'] => h4:fresh
p[style-name='Caption'] => p.caption:fresh
p[style-name='Code'] => pre:fresh
r[style-name='Code Char'] => code
"""


# ---------------------------------------------------------------------------
# Markdown cleanup
# ---------------------------------------------------------------------------

def _clean_markdown(md: str) -> str:
    """Remove unnecessary backslash escapes that mammoth adds."""
    # Convert mammoth bold markers (__text__) to standard markdown (**text**)
    md = re.sub(r"__(.+?)__", r"**\1**", md)
    return re.sub(r"\\([^`*\[\]()])", r"\1", md)


def _fully_unescape(text: str) -> str:
    """Fully strip all backslash escapes — used only for cell matching."""
    return re.sub(r"\\(.)", r"\1", text)


def _normalize_block(block: str) -> str:
    """Fully unescape and strip mammoth bold markers for cell comparison."""
    return _fully_unescape(block).strip().lstrip("_").rstrip("_").strip()


# ---------------------------------------------------------------------------
# Table rendering (python-docx)
# ---------------------------------------------------------------------------

def _cell_text(cell) -> str:
    """Extract plain text from a table cell."""
    return " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def _table_to_md(table) -> str:
    """Convert a python-docx Table to a markdown table string."""
    rows = table.rows
    if not rows:
        return ""

    grid = [[_cell_text(cell) for cell in row.cells] for row in rows]
    col_count = max(len(row) for row in grid)
    col_widths = [0] * col_count

    for row in grid:
        for i, cell in enumerate(row):
            col_widths[i] = max(col_widths[i], len(cell))

    def _fmt_row(cells: List[str]) -> str:
        padded = [
            cells[i].ljust(col_widths[i]) if i < len(cells) else " " * col_widths[i]
            for i in range(col_count)
        ]
        return "| " + " | ".join(padded) + " |"

    def _separator() -> str:
        return "| " + " | ".join("-" * w for w in col_widths) + " |"

    lines = [_fmt_row(grid[0]), _separator()]
    for row in grid[1:]:
        lines.append(_fmt_row(row))

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Cell set builder
# ---------------------------------------------------------------------------

def _build_cell_set(tables: list) -> Set[str]:
    """Collect all cell texts (fully unescaped) for block matching."""
    cell_texts: Set[str] = set()
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                text = _cell_text(cell).strip()
                if text:
                    cell_texts.add(text)
    return cell_texts


# ---------------------------------------------------------------------------
# Core: stitch mammoth blocks + python-docx tables in order
# ---------------------------------------------------------------------------

def _stitch(raw_blocks: List[str], md_tables: List[str], cell_texts: Set[str]) -> str:
    """
    Strategy:
    1. Walk raw_blocks and find the first index where a table-cell block appears.
       That is where the table belongs (right before that run of cell blocks).
    2. Insert the rendered markdown table at that position.
    3. Skip all subsequent cell blocks for that table.
    4. Repeat for each table.
    5. Clean escapes on all surviving non-cell blocks.
    """
    result: List[str] = []
    table_idx = 0
    i = 0

    while i < len(raw_blocks):
        block = raw_blocks[i]
        normalized = _normalize_block(block)

        if normalized in cell_texts:
            # This block is a table cell — insert the table once, then skip
            # all consecutive cell blocks
            if table_idx < len(md_tables):
                result.append(md_tables[table_idx])
                table_idx += 1
            # Skip all subsequent blocks that are also cell content
            while i < len(raw_blocks) and _normalize_block(raw_blocks[i]) in cell_texts:
                i += 1
        else:
            result.append(_clean_markdown(block))
            i += 1

    # Append any tables that had no matching cell blocks (e.g. empty tables)
    while table_idx < len(md_tables):
        result.append(md_tables[table_idx])
        table_idx += 1

    return "\n\n".join(b for b in result if b.strip()) + "\n"


# ---------------------------------------------------------------------------
# Core conversion
# ---------------------------------------------------------------------------

def convert_file(input_path: Path, output_path: Path, quiet: bool = False) -> bool:
    """Convert a single .docx file to markdown. Returns True on success."""
    try:
        # ── Step 1: mammoth — paragraphs, headings, lists ─────────────────────
        with open(input_path, "rb") as f:
            result = mammoth.convert_to_markdown(f, style_map=STYLE_MAP)

        if result.messages and not quiet:
            for msg in result.messages:
                print(f"  [warning] {msg}")

        # ── Step 2: python-docx — render tables ───────────────────────────────
        doc = DocxDocument(input_path)
        md_tables = [_table_to_md(t) for t in doc.tables]
        cell_texts = _build_cell_set(doc.tables)

        # ── Step 3: Split mammoth output and stitch ────────────────────────────
        raw_blocks = re.split(r"\n{2,}", result.value.strip())
        markdown = _stitch(raw_blocks, md_tables, cell_texts)

        # ── Step 4: Write output ──────────────────────────────────────────────
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(markdown, encoding="utf-8")

        if not quiet:
            print(f"  ✔  {input_path.name}  →  {output_path}")

        return True

    except Exception as e:
        print(f"  ✘  {input_path.name}: {e}", file=sys.stderr)
        return False


def process_folder(
    input_folder: Path,
    output_folder: Path,
    quiet: bool = False,
) -> dict:
    """Convert all .docx files in a folder. Returns a summary dict."""
    docx_files = sorted(input_folder.rglob("*.docx"))

    if not docx_files:
        print(f"No .docx files found in {input_folder}")
        return {"total": 0, "success": 0, "failed": 0}

    success, failed = 0, 0

    for docx_path in docx_files:
        relative = docx_path.relative_to(input_folder)
        out_path = (output_folder / relative).with_suffix(".md")

        if convert_file(docx_path, out_path, quiet):
            success += 1
        else:
            failed += 1

    summary = {"total": len(docx_files), "success": success, "failed": failed}

    if not quiet:
        print(
            f"\nDone — {success}/{len(docx_files)} converted"
            + (f", {failed} failed" if failed else "")
        )

    return summary


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Convert .docx files to markdown (mammoth + python-docx).",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )

    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("input", nargs="?", help="Path to a single .docx file")
    group.add_argument("--input-folder", type=Path, help="Folder containing .docx files")

    parser.add_argument("--output", type=Path, help="Output .md path (single-file mode)")
    parser.add_argument(
        "--output-folder",
        type=Path,
        default=Path("./markdown"),
        help="Output folder (folder mode, default: ./markdown)",
    )
    parser.add_argument("--quiet", action="store_true", help="Suppress progress output")

    return parser


def main():
    parser = build_parser()
    args = parser.parse_args()

    if args.input_folder:
        input_folder = args.input_folder
        if not input_folder.is_dir():
            parser.error(f"--input-folder does not exist: {input_folder}")

        summary = process_folder(input_folder, args.output_folder, args.quiet)
        sys.exit(0 if summary["failed"] == 0 else 1)

    input_path = Path(args.input)
    if not input_path.is_file():
        parser.error(f"File not found: {input_path}")
    if input_path.suffix.lower() != ".docx":
        parser.error(f"Expected a .docx file, got: {input_path.suffix}")

    output_path = args.output or input_path.with_suffix(".md")
    success = convert_file(input_path, output_path, args.quiet)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()